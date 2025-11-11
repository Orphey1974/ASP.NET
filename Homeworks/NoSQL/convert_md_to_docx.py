#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для конвертации Markdown файлов в Word формат (.docx).
Использует Python библиотеки markdown и python-docx.
Полная поддержка кириллицы и UTF-8.
Название файла остается тем же, меняется только расширение на .docx
"""

import os
import sys
from pathlib import Path
import html
import re

# Устанавливаем UTF-8 для Windows консоли
if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except (AttributeError, OSError):
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')


def check_dependencies():
    """Проверяет наличие необходимых библиотек"""
    try:
        import importlib
        importlib.import_module('markdown')
        importlib.import_module('docx')
        importlib.import_module('bs4')
        return True
    except ImportError as e:
        print(f"❌ Отсутствует необходимая библиотека: {e}")
        print("\n💡 Установите зависимости:")
        print("   pip install markdown python-docx beautifulsoup4")
        return False


def clean_text(text, remove_emoji=True):
    """Очищает текст от лишних пробелов и обрабатывает кодировку"""
    if not text:
        return ""

    if isinstance(text, str):
        text = text.strip()
    else:
        try:
            if isinstance(text, bytes):
                text = text.decode('utf-8').strip()
            else:
                text = str(text).strip()
        except (UnicodeDecodeError, UnicodeEncodeError):
            text = str(text).strip()

    text = ' '.join(text.split())
    text = html.unescape(text)

    replacements = {
        '&nbsp;': ' ',
        '&mdash;': '—',
        '&ndash;': '–',
        '&quot;': '"',
        '&amp;': '&',
        '&lt;': '<',
        '&gt;': '>',
        '&laquo;': '«',
        '&raquo;': '»',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    if remove_emoji:
        text = re.sub(r'[\U0001F300-\U0001F9FF]', '', text)
        text = re.sub(r'[\u2600-\u26FF]', '', text)
        text = re.sub(r'[\u2700-\u27BF]', '', text)

    try:
        import unicodedata
        text = unicodedata.normalize('NFKC', text)
    except (ImportError, AttributeError):
        try:
            text = text.encode('utf-8').decode('utf-8')
        except (UnicodeEncodeError, UnicodeDecodeError):
            pass

    return text.strip()


def convert_md_to_docx(md_file_path, output_dir=None):
    """Конвертирует MD файл в Word документ"""
    try:
        import markdown
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from bs4 import BeautifulSoup

        md_path = Path(md_file_path)

        if not md_path.exists():
            print(f"❌ Файл не найден: {md_file_path}")
            return False

        if not md_path.suffix.lower() == '.md':
            print(f"❌ Файл не является Markdown файлом: {md_file_path}")
            return False

        print(f"📄 Чтение файла: {md_path.name}")

        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        if isinstance(md_content, bytes):
            md_content = md_content.decode('utf-8')

        # Предобработка: удаляем пустые строки внутри таблиц
        # Таблицы в Markdown должны быть без пустых строк между строками
        lines = md_content.split('\n')
        processed_lines = []
        in_table = False

        for i, line in enumerate(lines):
            stripped = line.strip()
            # Проверяем, является ли строка частью таблицы
            is_table_separator = stripped.startswith('|') and '---' in stripped
            is_table_row = stripped.startswith('|') and not is_table_separator

            if is_table_row or is_table_separator:
                in_table = True
                processed_lines.append(line)
            elif in_table and stripped == '':
                # Пропускаем пустые строки внутри таблицы
                continue
            else:
                in_table = False
                processed_lines.append(line)

        md_content = '\n'.join(processed_lines)

        md = markdown.Markdown(
            extensions=[
                'extra',       # Дополнительные возможности (таблицы, fenced code blocks)
                'codehilite',  # Подсветка кода
                'tables',      # Таблицы (обязательно!)
                'toc',         # Оглавление
            ]
        )
        html_content = md.convert(md_content)

        if not isinstance(html_content, str):
            html_content = str(html_content, 'utf-8')

        if output_dir:
            output_path = Path(output_dir) / f"{md_path.stem}.docx"
        else:
            output_path = md_path.parent / f"{md_path.stem}.docx"

        print(f"📝 Конвертация в Word: {output_path.name}")

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        soup = BeautifulSoup(html_content, 'html.parser')

        def process_element(elem):
            """Рекурсивно обрабатывает HTML элементы"""
            if elem is None:
                return

            if elem.name in ['script', 'style']:
                return

            if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                text = clean_text(elem.get_text(strip=True, separator=' '))
                if text:
                    level = int(elem.name[1])
                    heading = doc.add_heading(text, level)
                    heading.style.font.name = 'Arial'

            elif elem.name == 'p':
                p = doc.add_paragraph()
                p.style.font.name = 'Arial'

                for content in elem.children:
                    if isinstance(content, str):
                        text = clean_text(content, remove_emoji=False)
                        if text:
                            run = p.add_run(text)
                            run.font.name = 'Arial'
                    elif hasattr(content, 'name'):
                        if content.name == 'strong' or content.name == 'b':
                            text = clean_text(content.get_text())
                            if text:
                                run = p.add_run(text)
                                run.font.bold = True
                                run.font.name = 'Arial'
                        elif content.name == 'em' or content.name == 'i':
                            text = clean_text(content.get_text())
                            if text:
                                run = p.add_run(text)
                                run.font.italic = True
                                run.font.name = 'Arial'
                        elif content.name == 'code':
                            text = clean_text(content.get_text())
                            if text:
                                run = p.add_run(text)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                        else:
                            text = clean_text(content.get_text())
                            if text:
                                run = p.add_run(text)
                                run.font.name = 'Arial'
                    elif hasattr(content, 'string') and content.string:
                        text = clean_text(str(content.string))
                        if text:
                            run = p.add_run(text)
                            run.font.name = 'Arial'

                if not p.runs:
                    if len(doc.paragraphs) > 0:
                        last_p = doc.paragraphs[-1]
                        if not last_p.runs:
                            last_p._element.getparent().remove(last_p._element)

            elif elem.name == 'pre':
                code_element = elem.find('code')
                is_mermaid = False
                diagram_code = None

                if code_element:
                    diagram_code = code_element.get_text()
                    class_attr = code_element.get('class', [])
                    if isinstance(class_attr, list):
                        class_attr = ' '.join(class_attr)
                    else:
                        class_attr = str(class_attr) if class_attr else ''

                    pre_text = elem.get_text().lower()

                    has_mermaid_class = (
                        'language-mermaid' in class_attr.lower() or
                        'lang-mermaid' in class_attr.lower() or
                        'mermaid' in class_attr.lower()
                    )
                    has_mermaid_syntax = (
                        'flowchart' in diagram_code.lower() or
                        'graph' in diagram_code.lower() or
                        'sequenceDiagram' in diagram_code.lower() or
                        'gantt' in diagram_code.lower() or
                        'pie' in diagram_code.lower() or
                        'stateDiagram' in diagram_code.lower()
                    )
                    has_mermaid_arrows = (
                        ('-->' in diagram_code and '[' in diagram_code
                         and ']' in diagram_code) or
                        ('--->' in diagram_code)
                    )
                    is_mermaid = (
                        has_mermaid_class or
                        has_mermaid_syntax or
                        has_mermaid_arrows or
                        'mermaid' in pre_text
                    )

                if is_mermaid and diagram_code:
                    # Добавляем заголовок для Mermaid диаграммы
                    heading = doc.add_paragraph("Диаграмма Mermaid")
                    heading.style.font.name = 'Arial'
                    heading.style.font.bold = True
                    heading.style.font.size = Pt(12)

                    clean_diagram_code = diagram_code.strip()
                    lines = clean_diagram_code.split('\n')
                    # Удаляем первую строку, если она содержит только "mermaid"
                    if lines and lines[0].lower().strip().startswith('mermaid'):
                        clean_diagram_code = '\n'.join(lines[1:]).strip()

                    # Удаляем пустые строки в начале и конце
                    clean_diagram_code = clean_diagram_code.strip()

                    # Пытаемся получить изображение диаграммы через API
                    image_added = False
                    try:
                        import base64
                        import requests
                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                        # Кодируем код диаграммы в base64 для API
                        diagram_base64 = base64.urlsafe_b64encode(
                            clean_diagram_code.encode('utf-8')).decode('utf-8')

                        # Используем mermaid.ink API для генерации изображения
                        api_url = f"https://mermaid.ink/img/{diagram_base64}"

                        print(f"   Попытка загрузки изображения диаграммы...")
                        response = requests.get(api_url, timeout=10)

                        if response.status_code == 200:
                            # Сохраняем изображение во временный файл
                            import tempfile
                            temp_img = tempfile.NamedTemporaryFile(
                                suffix='.png', delete=False)
                            temp_img.write(response.content)
                            temp_img.close()
                            temp_img_path = temp_img.name

                            # Добавляем изображение в документ
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run()
                            run.add_picture(temp_img_path, width=Inches(6))
                            image_added = True
                            print(f"   ✓ Изображение диаграммы добавлено")

                            # Удаляем временный файл
                            try:
                                os.unlink(temp_img_path)
                            except:
                                pass
                    except ImportError:
                        print(f"   ⚠️  Библиотека requests не установлена, пропускаю визуализацию")
                    except Exception as e:
                        print(f"   ⚠️  Не удалось загрузить изображение: {e}")

                    # Если изображение не удалось добавить, показываем код
                    if not image_added:
                        # Удаляем пустые строки в середине, но сохраняем структуру
                        code_lines = [line for line in clean_diagram_code.split('\n') if line.strip() or line == '']
                        # Если все строки пустые, используем оригинальный код
                        if not any(line.strip() for line in code_lines):
                            code_lines = clean_diagram_code.split('\n')

                        if code_lines:
                            # Создаем таблицу для отображения кода диаграммы
                            table = doc.add_table(rows=1, cols=1)
                            table.style = 'Light Grid Accent 1'
                            cell = table.rows[0].cells[0]

                            para_format = cell.paragraphs[0].paragraph_format
                            para_format.left_indent = Inches(0.1)
                            para_format.right_indent = Inches(0.1)
                            para_format.space_before = Pt(6)
                            para_format.space_after = Pt(6)

                            shading_elm = parse_xml(
                                r'<w:shd {} w:fill="F5F5F5"/>'.format(
                                    nsdecls('w')))
                            cell._element.get_or_add_tcPr().append(shading_elm)

                            # Добавляем код построчно
                            for i, line in enumerate(code_lines):
                                line_stripped = line.rstrip()

                                # Добавляем даже пустые строки для сохранения структуры
                                if i == 0:
                                    p = cell.paragraphs[0]
                                    if p.runs:
                                        p.clear()
                                    if line_stripped:
                                        run = p.add_run(line_stripped)
                                        run.font.name = 'Courier New'
                                        run.font.size = Pt(9)
                                else:
                                    p = cell.add_paragraph()
                                    if line_stripped:
                                        run = p.add_run(line_stripped)
                                        run.font.name = 'Courier New'
                                        run.font.size = Pt(9)

                            # Добавляем примечание
                            note_text = (
                                "Примечание: Для визуализации диаграммы используйте "
                                "редакторы с поддержкой Mermaid (VS Code с расширением, "
                                "draw.io, или онлайн-редакторы mermaid.live).")
                            note = doc.add_paragraph(note_text)
                            note.style.font.name = 'Arial'
                            note.style.font.italic = True
                            note.style.font.size = Pt(9)
                            note.style.font.color.rgb = RGBColor(128, 128, 128)

                    # Добавляем отступ после диаграммы
                    doc.add_paragraph()
                else:
                    code_text = elem.get_text()
                    if code_text:
                        code_lines = code_text.split('\n')

                        table = doc.add_table(rows=1, cols=1)
                        table.style = 'Light Grid Accent 1'
                        cell = table.rows[0].cells[0]

                        para_format = cell.paragraphs[0].paragraph_format
                        para_format.left_indent = Inches(0.1)
                        para_format.right_indent = Inches(0.1)
                        para_format.space_before = Pt(6)
                        para_format.space_after = Pt(6)

                        shading_elm = parse_xml(
                            r'<w:shd {} w:fill="F5F5F5"/>'.format(
                                nsdecls('w')))
                        cell._element.get_or_add_tcPr().append(shading_elm)

                        for i, line in enumerate(code_lines):
                            line_clean = clean_text(line, remove_emoji=False)
                            if line_clean or i == 0:
                                if i == 0:
                                    p = cell.paragraphs[0]
                                    run = p.add_run(line_clean)
                                else:
                                    p = cell.add_paragraph()
                                    run = p.add_run(line_clean)

                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)

            elif elem.name in ['ul', 'ol']:
                for li in elem.find_all('li', recursive=False):
                    text = clean_text(li.get_text())
                    if text:
                        if elem.name == 'ol':
                            p = doc.add_paragraph(text, style='List Number')
                        else:
                            p = doc.add_paragraph(text, style='List Bullet')
                        p.style.font.name = 'Arial'

            elif elem.name == 'table':
                table_data = []
                max_cols = 0

                # Собираем данные таблицы
                for tr in elem.find_all('tr'):
                    row = []
                    for td in tr.find_all(['td', 'th']):
                        # Получаем текст из ячейки, сохраняя структуру
                        cell_text = td.get_text(separator=' ', strip=True)
                        # Очищаем от лишних пробелов, но сохраняем содержимое
                        cell_text = ' '.join(cell_text.split())
                        row.append(cell_text)

                    if row:
                        # Нормализуем количество колонок
                        max_cols = max(max_cols, len(row))
                        table_data.append(row)

                if table_data:
                    # Нормализуем все строки до одинакового количества колонок
                    for row in table_data:
                        while len(row) < max_cols:
                            row.append('')

                    # Создаем таблицу в Word
                    table = doc.add_table(rows=len(table_data),
                                          cols=max_cols)
                    table.style = 'Light Grid Accent 1'

                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            if j < max_cols:
                                cell = table.rows[i].cells[j]
                                # Очищаем ячейку перед добавлением текста
                                cell.text = ''
                                # Добавляем текст с правильным форматированием
                                p = cell.paragraphs[0]
                                run = p.add_run(cell_text)

                                # Заголовок таблицы (первая строка) - жирный
                                if i == 0:
                                    run.font.bold = True
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                else:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)

                                # Настройка выравнивания
                                p.alignment = 1  # Left alignment

                                # Настройка отступов в ячейке
                                cell.vertical_alignment = 1  # Top alignment

                    # Добавляем отступ после таблицы
                    doc.add_paragraph()

            elif elem.name == 'hr':
                pass

        body = soup.find('body')
        root = body if body else soup

        block_tags = [
            'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
            'p', 'ul', 'ol', 'table', 'pre',
            'blockquote', 'hr'
        ]

        # Сначала обрабатываем все таблицы отдельно
        all_tables = root.find_all('table')
        processed_table_ids = set()
        print(f"   Найдено таблиц в HTML: {len(all_tables)}")
        for table_elem in all_tables:
            # Проверяем, что таблица не вложена в другую таблицу
            is_nested = False
            for parent in table_elem.parents:
                if parent.name == 'table':
                    is_nested = True
                    break
            if not is_nested:
                table_id = id(table_elem)
                if table_id not in processed_table_ids:
                    processed_table_ids.add(table_id)
                    process_element(table_elem)

        # Сначала обрабатываем все Mermaid диаграммы отдельно
        all_pre_blocks = root.find_all('pre')
        processed_pre_ids = set()
        mermaid_count = 0

        for pre_elem in all_pre_blocks:
            code_elem = pre_elem.find('code')
            if code_elem:
                code_text = code_elem.get_text()
                class_attr = code_elem.get('class', [])
                if isinstance(class_attr, list):
                    class_attr = ' '.join(class_attr)
                else:
                    class_attr = str(class_attr) if class_attr else ''

                pre_text = pre_elem.get_text().lower()

                has_mermaid_class = (
                    'language-mermaid' in class_attr.lower() or
                    'lang-mermaid' in class_attr.lower() or
                    'mermaid' in class_attr.lower()
                )
                has_mermaid_syntax = (
                    'flowchart' in code_text.lower() or
                    'graph' in code_text.lower() or
                    'sequencediagram' in code_text.lower() or
                    'gantt' in code_text.lower() or
                    'pie' in code_text.lower() or
                    'statediagram' in code_text.lower()
                )
                has_mermaid_arrows = (
                    ('-->' in code_text and '[' in code_text and ']' in code_text) or
                    ('--->' in code_text)
                )
                is_mermaid = (
                    has_mermaid_class or
                    has_mermaid_syntax or
                    has_mermaid_arrows or
                    'mermaid' in pre_text
                )

                if is_mermaid:
                    mermaid_count += 1
                    # Проверяем, что блок не вложен в другой элемент
                    is_nested = False
                    for parent in pre_elem.parents:
                        if parent.name in ['td', 'th', 'li', 'table']:
                            is_nested = True
                            break

                    if not is_nested:
                        pre_id = id(pre_elem)
                        if pre_id not in processed_pre_ids:
                            processed_pre_ids.add(pre_id)
                            print(f"   Обработка Mermaid диаграммы #{mermaid_count}...")
                            process_element(pre_elem)
                            print(f"   ✓ Mermaid диаграмма #{mermaid_count} добавлена в документ")

        print(f"   Найдено Mermaid диаграмм: {mermaid_count}, обработано: {len(processed_pre_ids)}")

        # Затем обрабатываем остальные элементы
        all_elements = root.find_all(block_tags)

        processed_ids = set()

        for elem in all_elements:
            # Пропускаем таблицы, так как они уже обработаны
            if elem.name == 'table':
                continue
            # Пропускаем блоки pre с Mermaid, так как они уже обработаны
            if elem.name == 'pre':
                pre_id = id(elem)
                if pre_id in processed_pre_ids:
                    continue

            parent = elem.parent
            if parent:
                # Пропускаем элементы внутри таблиц (они обрабатываются вместе с таблицей)
                if parent.name in ['td', 'th', 'thead', 'tbody', 'tr', 'table']:
                    continue
                # Пропускаем элементы внутри списков (обрабатываются вместе со списком)
                if parent.name in ['li']:
                    continue
                parents_list = list(elem.parents)
                if len(parents_list) > 3:
                    continue
                parent_names = ['body', 'html', '[document]', 'div']
                if (parent.name not in parent_names and
                        len(parents_list) > 2):
                    continue

            elem_id = id(elem)
            if elem_id not in processed_ids:
                processed_ids.add(elem_id)
                process_element(elem)

        doc.save(str(output_path))

        file_size = output_path.stat().st_size
        file_size_kb = file_size / 1024

        print("✅ Word документ успешно создан!")
        print(f"   Путь: {output_path}")
        print(f"   Размер: {file_size_kb:.2f} KB")

        return True

    except Exception as e:
        print(f"❌ Ошибка при конвертации: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Основная функция"""
    print("=" * 60)
    print("📄 КОНВЕРТАЦИЯ MARKDOWN В WORD")
    print("=" * 60)
    print()

    if not check_dependencies():
        print("\n❌ Установите необходимые библиотеки перед использованием")
        return 1

    if len(sys.argv) < 2:
        print("❌ Укажите путь к Markdown файлу")
        print("\n💡 Использование:")
        print(f"   python {sys.argv[0]} <путь_к_файлу.md>")
        return 1

    md_file = sys.argv[1]

    if convert_md_to_docx(md_file):
        print("\n✅ Конвертация завершена успешно!")
        return 0
    else:
        print("\n❌ Произошла ошибка при конвертации.")
        return 1


if __name__ == '__main__':
    exit_code = main()
    is_debugging = (
        'pydevd' in sys.modules or
        'debugpy' in sys.modules or
        os.getenv('VSCODE_INSPECTOR_OPTIONS') is not None
    )

    if not is_debugging and exit_code != 0:
        sys.exit(exit_code)

